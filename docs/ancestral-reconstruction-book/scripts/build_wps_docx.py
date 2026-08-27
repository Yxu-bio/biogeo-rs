from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips


INK = "0B2545"
HEADING = "2E74B5"
HEADING_DARK = "1F4D78"
MUTED = "6B7280"
TABLE_HEADER = "E8EEF5"
CALLOUT_FILL = "F4F6F9"
LATIN_FONT = "Calibri"
EAST_ASIA_BODY = "Microsoft YaHei"
EAST_ASIA_HEADING = "Microsoft YaHei"
MATH_FONT = "Cambria Math"
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_fonts(element, latin: str, east_asia: str) -> None:
    r_pr = element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), latin)
    r_fonts.set(qn("w:hAnsi"), latin)
    r_fonts.set(qn("w:eastAsia"), east_asia)
    r_fonts.set(qn("w:cs"), latin)


def set_style(
    style,
    *,
    size: float,
    color: str = "000000",
    bold: bool = False,
    before: float = 0,
    after: float = 6,
    line_spacing: float = 1.25,
    latin: str = LATIN_FONT,
    east_asia: str = EAST_ASIA_BODY,
) -> None:
    style.font.name = latin
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)
    set_fonts(style.element, latin, east_asia)
    paragraph = style.paragraph_format
    paragraph.space_before = Pt(before)
    paragraph.space_after = Pt(after)
    paragraph.line_spacing = line_spacing
    paragraph.widow_control = True


def get_or_add(parent, tag: str):
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.append(child)
    return child


def shade_paragraph(paragraph, fill: str, border_color: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shading = get_or_add(p_pr, "w:shd")
    shading.set(qn("w:fill"), fill)
    borders = get_or_add(p_pr, "w:pBdr")
    left = get_or_add(borders, "w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), border_color)


def set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def keep_row_together(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = get_or_add(tc_pr, "w:tcMar")
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = get_or_add(margins, f"w:{edge}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width: int) -> None:
    cell.width = Twips(width)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = get_or_add(tc_pr, "w:tcW")
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = get_or_add(tbl_pr, "w:tblW")
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = get_or_add(tbl_pr, "w:tblInd")
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = get_or_add(tbl_pr, "w:tblLayout")
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row_index, row in enumerate(table.rows):
        keep_row_together(row)
        for column_index, cell in enumerate(row.cells):
            width = widths[min(column_index, len(widths) - 1)]
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index == 0:
                shading = get_or_add(cell._tc.get_or_add_tcPr(), "w:shd")
                shading.set(qn("w:fill"), TABLE_HEADER)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.15
                for run in paragraph.runs:
                    set_fonts(run._element, LATIN_FONT, EAST_ASIA_BODY)
                    run.font.size = Pt(9.5)
                    if row_index == 0:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor.from_string(INK)
        if row_index == 0:
            set_repeat_header(row)


def widths_for(table) -> list[int]:
    headers = tuple(cell.text.strip() for cell in table.rows[0].cells)
    patterns = {
        ("你的目标", "建议先读", "暂时可以后读"): [2700, 3330, 3330],
        ("符号", "本书中的含义", "是观测、隐藏量还是参数"): [1600, 5000, 2760],
        ("末端", "观测状态"): [2500, 6860],
        ("事件", "一阶概率", "对“没有采样后代”概率的后续贡献"): [1800, 2200, 5360],
        ("模型", "容易互相补偿的量", "表面现象"): [1500, 3300, 4560],
        ("研究问题", "输入", "数学核心", "主要输出", "最常见误读"): [1500, 1700, 2300, 1800, 2060],
        ("情况", "更合理的基线", "原因"): [2100, 3000, 4260],
    }
    if headers in patterns:
        return patterns[headers]
    count = len(table.rows[0].cells)
    base = CONTENT_WIDTH_DXA // count
    widths = [base] * count
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    return widths


def add_page_field(paragraph) -> None:
    paragraph.add_run("第 ")
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, value, end))
    paragraph.add_run(" 页")


def configure_settings(document: Document) -> None:
    settings = document.settings.element
    update = get_or_add(settings, "w:updateFields")
    update.set(qn("w:val"), "true")
    language = get_or_add(settings, "w:themeFontLang")
    language.set(qn("w:val"), "zh-CN")
    language.set(qn("w:eastAsia"), "zh-CN")
    get_or_add(settings, "w:doNotUseHTMLParagraphAutoSpacing")

    math_pr = settings.find(qn("m:mathPr"))
    if math_pr is None:
        math_pr = OxmlElement("m:mathPr")
        settings.append(math_pr)
    math_font = math_pr.find(qn("m:mathFont"))
    if math_font is None:
        math_font = OxmlElement("m:mathFont")
        math_pr.append(math_font)
    math_font.set(qn("m:val"), MATH_FONT)


def style_document(input_path: Path, output_path: Path) -> None:
    document = Document(input_path)
    section = document.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    styles = document.styles
    set_style(styles["Normal"], size=11)
    for name in ("Body Text", "First Paragraph"):
        if name in styles:
            set_style(styles[name], size=11)
    if "Compact" in styles:
        set_style(styles["Compact"], size=11, after=4)
    if "Block Text" in styles:
        set_style(styles["Block Text"], size=10.5, color=INK, after=8)
    set_style(styles["Title"], size=24, color=INK, bold=True, before=120, after=12, line_spacing=1.1, east_asia=EAST_ASIA_HEADING)
    set_style(styles["Subtitle"], size=14, color=HEADING_DARK, before=0, after=10, line_spacing=1.1, east_asia=EAST_ASIA_HEADING)
    if "Date" in styles:
        set_style(styles["Date"], size=10, color=MUTED, before=4, after=0, line_spacing=1.0)
    set_style(styles["Heading 1"], size=16, color=HEADING, bold=True, before=18, after=10, east_asia=EAST_ASIA_HEADING)
    set_style(styles["Heading 2"], size=13, color=HEADING, bold=True, before=14, after=7, east_asia=EAST_ASIA_HEADING)
    set_style(styles["Heading 3"], size=12, color=HEADING_DARK, bold=True, before=10, after=5, east_asia=EAST_ASIA_HEADING)
    for name, size, indent in (("TOC 1", 11, 0), ("TOC 2", 10.5, 0.22), ("TOC 3", 10, 0.44)):
        if name in styles:
            set_style(styles[name], size=size, color=INK, after=3, line_spacing=1.1)
            styles[name].paragraph_format.left_indent = Inches(indent)

    date_seen = False
    heading_one_seen = False
    for paragraph in document.paragraphs:
        style_name = paragraph.style.name if paragraph.style is not None else ""
        if style_name == "Title":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif style_name in ("Subtitle", "Date"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if style_name == "Date" and not date_seen:
            paragraph.add_run().add_break(WD_BREAK.PAGE)
            date_seen = True
        if style_name == "Heading 1":
            paragraph.paragraph_format.keep_with_next = True
            if not heading_one_seen:
                paragraph.paragraph_format.page_break_before = True
                heading_one_seen = True
        elif style_name in ("Heading 2", "Heading 3"):
            paragraph.paragraph_format.keep_with_next = True
        if style_name == "Block Text":
            paragraph.paragraph_format.left_indent = Inches(0.18)
            paragraph.paragraph_format.right_indent = Inches(0.12)
            paragraph.paragraph_format.space_before = Pt(6)
            shade_paragraph(paragraph, CALLOUT_FILL, HEADING)
        if paragraph._p.xpath(".//m:oMath"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(5)
            paragraph.paragraph_format.space_after = Pt(7)
            paragraph.paragraph_format.keep_together = True
        if paragraph._p.pPr is not None and paragraph._p.pPr.numPr is not None:
            paragraph.paragraph_format.left_indent = Inches(0.375)
            paragraph.paragraph_format.first_line_indent = Inches(-0.188)
            paragraph.paragraph_format.space_after = Pt(4)
            paragraph.paragraph_format.line_spacing = 1.25
        for run in paragraph.runs:
            heading_font = EAST_ASIA_HEADING if style_name.startswith("Heading") or style_name in ("Title", "Subtitle") else EAST_ASIA_BODY
            set_fonts(run._element, LATIN_FONT, heading_font)

    for table in document.tables:
        set_table_geometry(table, widths_for(table))

    header = section.header
    header_paragraph = header.paragraphs[0]
    header_paragraph.text = "祖先重建的数学"
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_paragraph.paragraph_format.space_after = Pt(0)
    for run in header_paragraph.runs:
        set_fonts(run._element, LATIN_FONT, EAST_ASIA_HEADING)
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor.from_string(MUTED)

    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.clear()
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_field(footer_paragraph)
    for run in footer_paragraph.runs:
        set_fonts(run._element, LATIN_FONT, EAST_ASIA_BODY)
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor.from_string(MUTED)

    configure_settings(document)
    properties = document.core_properties
    properties.title = "祖先重建的数学"
    properties.subject = "概率、树上剪枝、DEC、SSE 与模拟训练"
    properties.author = "BGB Rust 项目"
    properties.keywords = "祖先重建, 历史生物地理, DEC, SSE, 数学"
    document.save(output_path)


def main() -> None:
    if len(sys.argv) != 3:
        raise SystemExit("Usage: build_wps_docx.py INPUT.docx OUTPUT.docx")
    input_path = Path(sys.argv[1]).resolve()
    output_path = Path(sys.argv[2]).resolve()
    style_document(input_path, output_path)
    print(output_path)


if __name__ == "__main__":
    main()
